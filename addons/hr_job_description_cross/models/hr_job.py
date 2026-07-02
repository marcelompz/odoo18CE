# -*- coding: utf-8 -*-
"""Extension de hr.job para Manual de Funciones con workflow de aprobacion."""
import base64
import io

from odoo import _, api, fields, models
from odoo.exceptions import UserError

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class HrJob(models.Model):
    _inherit = 'hr.job'

    # === Codigo unico del manual (generado por secuencia) ===
    manual_code = fields.Char(
        string='Codigo del Manual',
        readonly=True, copy=False, index=True,
        help='Codigo unico del documento (auto-generado: MF/AAAA/NNNNN).',
    )

    # === I. Informacion Basica ===
    manual_revision = fields.Char(
        string='Revision', default='00', readonly=True, copy=False, tracking=True,
        help='Numero de revision del manual. Se autoincrementa cada vez que el manual aprobado vuelve a borrador.',
    )
    manual_boss_job_id = fields.Many2one('hr.job', string='Jefe Inmediato (Puesto)')
    manual_supervises_job_ids = fields.Many2many(
        'hr.job', relation='hr_job_supervises_rel',
        column1='job_id', column2='supervised_job_id',
        string='Puestos que Supervisa',
    )

    # === II. Naturaleza ===
    manual_nature = fields.Html(string='Naturaleza del Puesto')

    # === III. Funciones ===
    manual_function_line_ids = fields.One2many(
        'hr.job.function.line', 'job_id',
        string='Funciones y Responsabilidades',
    )

    # === IV. Requisitos ===
    manual_experience = fields.Html(string='Experiencia Requerida')
    manual_education_ids = fields.Many2many(
        'hr.education.level',
        relation='hr_job_education_rel',
        column1='job_id', column2='education_id',
        string='Formacion Requerida',
    )
    manual_education_notes = fields.Html(string='Formacion - Notas Adicionales')
    manual_skill_ids = fields.Many2many(
        'hr.skill',
        relation='hr_job_required_skill_rel',
        column1='job_id', column2='skill_id',
        string='Habilidades Requeridas',
    )
    manual_skills_notes = fields.Html(string='Habilidades - Notas Adicionales')
    manual_epi = fields.Html(string='EPI (Elementos de Proteccion Individual)')

    # === WORKFLOW ===
    manual_state = fields.Selection([
        ('draft', 'Borrador'),
        ('review', 'En Revision'),
        ('approved', 'Aprobado'),
        ('rejected', 'Rechazado'),
    ], string='Estado del Manual', default='draft', tracking=True,
        copy=False, index=True)

    manual_create_date = fields.Datetime(
        string='Fecha de Creacion', default=fields.Datetime.now,
        readonly=True, copy=False,
    )
    manual_creator_id = fields.Many2one(
        'res.users', string='Creador del Documento',
        default=lambda self: self.env.user, readonly=True, copy=False,
    )
    manual_review_date = fields.Datetime(
        string='Fecha de Envio a Revision', readonly=True, copy=False,
    )
    manual_approval_date = fields.Datetime(
        string='Fecha de Aprobacion', readonly=True, copy=False,
    )
    manual_approved_by_id = fields.Many2one(
        'res.users', string='Aprobado por', readonly=True, copy=False, tracking=True,
    )
    manual_rejection_reason = fields.Text(string='Motivo de Rechazo', copy=False)

    manual_next_revision_date = fields.Date(
        string='Proxima Revision', compute='_compute_manual_next_revision',
        store=True, copy=False,
    )
    manual_is_overdue = fields.Boolean(
        string='Vencido', compute='_compute_manual_is_overdue',
        search='_search_manual_is_overdue', store=False,
    )

    manual_can_approve = fields.Boolean(
        string='Puede aprobar', compute='_compute_manual_can_approve',
    )
    manual_is_locked = fields.Boolean(
        string='Bloqueado por aprobacion',
        compute='_compute_manual_is_locked', store=False,
    )

    @api.depends('manual_create_date')
    def _compute_manual_next_revision(self):
        from dateutil.relativedelta import relativedelta
        for rec in self:
            if rec.manual_create_date:
                rec.manual_next_revision_date = (
                    rec.manual_create_date.date() + relativedelta(months=3)
                )
            else:
                rec.manual_next_revision_date = False

    @api.depends('manual_next_revision_date')
    def _compute_manual_is_overdue(self):
        today = fields.Date.context_today(self)
        for rec in self:
            rec.manual_is_overdue = bool(
                rec.manual_next_revision_date
                and rec.manual_next_revision_date < today
            )

    def _search_manual_is_overdue(self, operator, value):
        today = fields.Date.context_today(self)
        if operator not in ('=', '!='):
            return []
        is_true = bool(value) if operator == '=' else not bool(value)
        if is_true:
            # Vencidos: next_revision_date existe y es < today
            return [('manual_next_revision_date', '<', today)]
        # Vigentes: sin fecha o fecha >= today
        return ['|',
                ('manual_next_revision_date', '=', False),
                ('manual_next_revision_date', '>=', today)]

    @api.depends_context('uid')
    def _compute_manual_can_approve(self):
        is_approver = self.env.user.has_group(
            'hr_job_description_cross.group_hr_job_manual_approver'
        ) or self.env.user.has_group('hr.group_hr_manager')
        for rec in self:
            rec.manual_can_approve = is_approver

    @api.depends('manual_state')
    def _compute_manual_is_locked(self):
        for rec in self:
            rec.manual_is_locked = rec.manual_state == 'approved'

    @api.model_create_multi
    def create(self, vals_list):
        for vals in vals_list:
            if not vals.get('manual_code'):
                vals['manual_code'] = self.env['ir.sequence'].next_by_code(
                    'hr.job.manual.code'
                ) or '/'
        return super().create(vals_list)

    # === Acciones de transicion ===
    def action_send_for_review(self):
        for rec in self:
            if rec.manual_state not in ('draft', 'rejected'):
                raise UserError(_(
                    'Solo se puede enviar a revision desde estado Borrador o Rechazado.'
                ))
            rec.write({
                'manual_state': 'review',
                'manual_review_date': fields.Datetime.now(),
                'manual_rejection_reason': False,
            })
        return True

    def action_approve(self):
        if not self.env.user.has_group(
            'hr_job_description_cross.group_hr_job_manual_approver'
        ) and not self.env.user.has_group('hr.group_hr_manager'):
            raise UserError(_(
                'Solo un usuario con el grupo "Aprobador de Manual de Funciones" puede aprobar.'
            ))
        for rec in self:
            if rec.manual_state != 'review':
                raise UserError(_('Solo se puede aprobar manuales en estado "En Revision".'))
            rec.write({
                'manual_state': 'approved',
                'manual_approval_date': fields.Datetime.now(),
                'manual_approved_by_id': self.env.user.id,
            })
        return True

    def action_reject(self):
        if not self.env.user.has_group(
            'hr_job_description_cross.group_hr_job_manual_approver'
        ) and not self.env.user.has_group('hr.group_hr_manager'):
            raise UserError(_('Solo un aprobador puede rechazar.'))
        for rec in self:
            if rec.manual_state != 'review':
                raise UserError(_('Solo se puede rechazar manuales en estado "En Revision".'))
            rec.write({'manual_state': 'rejected'})
        return True

    def action_reset_to_draft(self):
        if not self.env.user.has_group(
            'hr_job_description_cross.group_hr_job_manual_approver'
        ) and not self.env.user.has_group('hr.group_hr_manager'):
            raise UserError(_('Solo un aprobador puede reabrir un manual aprobado.'))
        for rec in self:
            new_rev = rec.manual_revision or '00'
            try:
                new_rev = '{:02d}'.format(int(new_rev) + 1)
            except (ValueError, TypeError):
                new_rev = '01'
            rec.write({
                'manual_state': 'draft',
                'manual_approval_date': False,
                'manual_approved_by_id': False,
                'manual_revision': new_rev,
                'manual_create_date': fields.Datetime.now(),
            })
        return True

    # Campos de CONTENIDO del manual que quedan bloqueados al aprobar
    # (no incluye computed/auto-mantenidos como next_revision_date)
    _MANUAL_PROTECTED_FIELDS = {
        'manual_boss_job_id', 'manual_supervises_job_ids', 'manual_nature',
        'manual_function_line_ids', 'manual_experience',
        'manual_education_ids', 'manual_education_notes',
        'manual_skill_ids', 'manual_skills_notes', 'manual_epi',
    }

    def write(self, vals):
        # Bloqueamos solo si tocan campos de contenido editable Y el manual
        # esta aprobado. Migraciones, recomputes y transiciones de estado
        # nunca quedan bloqueados.
        if self._MANUAL_PROTECTED_FIELDS & set(vals.keys()):
            # Si el contexto explicitamente desactiva el lock (migraciones)
            if not self.env.context.get('skip_manual_lock'):
                for rec in self:
                    if rec.manual_state == 'approved':
                        raise UserError(_(
                            'El manual del puesto "%s" esta APROBADO y bloqueado. '
                            'Para modificarlo, primero use "Reabrir (Borrador)" '
                            '(solo disponible para aprobadores).'
                        ) % rec.name)
        return super().write(vals)

    def action_print_job_manual(self):
        self.ensure_one()
        return self.env.ref(
            'hr_job_description_cross.action_report_hr_job_manual'
        ).report_action(self)

    # === EXPORTAR A WORD (.docx) ===
    def action_export_manual_docx(self):
        self.ensure_one()
        if not HAS_DOCX:
            raise UserError(_(
                'La libreria python-docx no esta instalada. '
                'Ejecute: pip install python-docx --break-system-packages'
            ))
        doc = self._build_docx_manual()
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = 'Manual_Funciones_%s_%s.docx' % (
            (self.manual_code or 'sn').replace('/', '_'),
            (self.name or 'puesto').replace(' ', '_'),
        )
        attachment = self.env['ir.attachment'].create({
            'name': filename,
            'datas': base64.b64encode(buffer.getvalue()),
            'res_model': 'hr.job',
            'res_id': self.id,
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        })
        return {
            'type': 'ir.actions.act_url',
            'url': '/web/content/%s?download=true' % attachment.id,
            'target': 'self',
        }

    def _strip_html(self, html_text):
        if not html_text:
            return ''
        try:
            from lxml import html as lxml_html
            return lxml_html.fromstring(html_text).text_content().strip()
        except Exception:
            import re
            return re.sub(r'<[^>]+>', '', html_text).strip()

    def _build_docx_manual(self):
        """Construye un documento .docx con el manual de funciones."""
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.8)

        # Cabecera: logo + titulo + revision
        header_table = doc.add_table(rows=1, cols=3)
        header_table.style = 'Table Grid'
        cell_logo = header_table.rows[0].cells[0]
        cell_title = header_table.rows[0].cells[1]
        cell_rev = header_table.rows[0].cells[2]

        if self.company_id.logo:
            try:
                logo_buf = io.BytesIO(base64.b64decode(self.company_id.logo))
                cell_logo.paragraphs[0].add_run().add_picture(logo_buf, width=Inches(1.2))
            except Exception:
                pass

        title_para = cell_title.paragraphs[0]
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run('DESCRIPCION DE FUNCIONES')
        title_run.bold = True
        title_run.font.size = Pt(14)
        cell_title.add_paragraph(self.company_id.name or '').alignment = WD_ALIGN_PARAGRAPH.CENTER

        rev_para = cell_rev.paragraphs[0]
        rev_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rev_run = rev_para.add_run('Codigo: %s\nRevision: %s' % (
            self.manual_code or '-', self.manual_revision or '00'))
        rev_run.bold = True

        # === METADATA: codigo, fechas, responsables ===
        def _fmt_date(dt, with_time=False):
            if not dt:
                return '-'
            if with_time:
                return dt.strftime('%d/%m/%Y %H:%M')
            return dt.strftime('%d/%m/%Y')

        meta_table = doc.add_table(rows=4, cols=4)
        meta_table.style = 'Table Grid'
        meta_table.rows[0].cells[0].text = 'Codigo del Manual'
        meta_table.rows[0].cells[1].text = self.manual_code or '-'
        meta_table.rows[0].cells[2].text = 'Estado'
        state_label = dict(self._fields['manual_state'].selection).get(
            self.manual_state, self.manual_state or '-')
        meta_table.rows[0].cells[3].text = state_label
        meta_table.rows[1].cells[0].text = 'Fecha de Creacion'
        meta_table.rows[1].cells[1].text = _fmt_date(self.manual_create_date, True)
        meta_table.rows[1].cells[2].text = 'Creador del Documento'
        meta_table.rows[1].cells[3].text = (
            self.manual_creator_id.name if self.manual_creator_id else '-')
        meta_table.rows[2].cells[0].text = 'Fecha Envio a Revision'
        meta_table.rows[2].cells[1].text = _fmt_date(self.manual_review_date, True)
        meta_table.rows[2].cells[2].text = 'Aprobador'
        meta_table.rows[2].cells[3].text = (
            self.manual_approved_by_id.name if self.manual_approved_by_id else '-')
        meta_table.rows[3].cells[0].text = 'Proxima Revision'
        next_rev_txt = _fmt_date(self.manual_next_revision_date)
        if self.manual_is_overdue:
            next_rev_txt += '  [VENCIDO]'
        meta_table.rows[3].cells[1].text = next_rev_txt
        meta_table.rows[3].cells[2].text = 'Fecha de Aprobacion'
        meta_table.rows[3].cells[3].text = _fmt_date(self.manual_approval_date, True)

        # Negrita en labels y color rojo si vencido
        for row in meta_table.rows:
            for col_idx in (0, 2):
                for p in row.cells[col_idx].paragraphs:
                    for r in p.runs:
                        r.bold = True
        if self.manual_is_overdue:
            for p in meta_table.rows[3].cells[1].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        # I. Informacion Basica
        doc.add_paragraph()
        h1 = doc.add_paragraph()
        h1_run = h1.add_run('I. INFORMACION BASICA')
        h1_run.bold = True
        h1_run.font.size = Pt(12)

        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'
        info_table.rows[0].cells[0].text = 'CARGO'
        info_table.rows[0].cells[1].text = self.name or ''
        info_table.rows[1].cells[0].text = 'JEFE'
        info_table.rows[1].cells[1].text = self.manual_boss_job_id.name if self.manual_boss_job_id else '-'
        info_table.rows[2].cells[0].text = 'SUPERVISA A'
        info_table.rows[2].cells[1].text = ', '.join(self.manual_supervises_job_ids.mapped('name')) or '-'

        # II. Naturaleza
        doc.add_paragraph()
        h2 = doc.add_paragraph()
        h2_run = h2.add_run('II. NATURALEZA DEL PUESTO')
        h2_run.bold = True
        h2_run.font.size = Pt(12)
        doc.add_paragraph(self._strip_html(self.manual_nature) or '-')

        # III. Funciones
        doc.add_paragraph()
        h3 = doc.add_paragraph()
        h3_run = h3.add_run('III. FUNCIONES Y RESPONSABILIDADES')
        h3_run.bold = True
        h3_run.font.size = Pt(12)
        for line in self.manual_function_line_ids.sorted('sequence'):
            p = doc.add_paragraph(style='List Number')
            run_name = p.add_run((line.name or (line.function_id.name if line.function_id else '')) + ': ')
            run_name.bold = True
            p.add_run(self._strip_html(line.description) or '')

        # IV. Requisitos
        doc.add_paragraph()
        h4 = doc.add_paragraph()
        h4_run = h4.add_run('IV. REQUISITOS MINIMOS PARA EL PUESTO')
        h4_run.bold = True
        h4_run.font.size = Pt(12)

        req_table = doc.add_table(rows=4, cols=2)
        req_table.style = 'Table Grid'
        req_table.rows[0].cells[0].text = 'EXPERIENCIA'
        req_table.rows[0].cells[1].text = self._strip_html(self.manual_experience) or '-'
        req_table.rows[1].cells[0].text = 'FORMACION'
        edu_list = ', '.join(self.manual_education_ids.mapped('name'))
        edu_notes = self._strip_html(self.manual_education_notes)
        req_table.rows[1].cells[1].text = '\n'.join(filter(None, [edu_list, edu_notes])) or '-'
        req_table.rows[2].cells[0].text = 'HABILIDAD'
        skill_list = ', '.join(self.manual_skill_ids.mapped('name'))
        skill_notes = self._strip_html(self.manual_skills_notes)
        req_table.rows[2].cells[1].text = '\n'.join(filter(None, [skill_list, skill_notes])) or '-'
        req_table.rows[3].cells[0].text = 'EPI'
        req_table.rows[3].cells[1].text = self._strip_html(self.manual_epi) or '-'

        # Firmas
        doc.add_paragraph()
        doc.add_paragraph()
        sign_table = doc.add_table(rows=1, cols=2)
        sign_table.rows[0].cells[0].text = '\n\n______________________\nElaborado por\nRecursos Humanos'
        sign_table.rows[0].cells[1].text = '\n\n______________________\nAprobado por\nDireccion'

        return doc
